from fastapi import FastAPI, UploadFile, File
from resume_parser import parse_resume, parse_job_description
from scorer import compute_match_score
from tips_generator import generate_tips


app = FastAPI()

@app.post("/analyze")
async def analyze(resume: UploadFile = File(...), jd: UploadFile = File(...)):
    resume_text = await resume.read()
    jd_text = await jd.read()

    resume_str = parse_resume(resume_text)
    jd_str = parse_job_description(jd_text)

    score = compute_match_score(resume_str, jd_str)
    tips = generate_tips(resume_str, jd_str)

    return {"match_score": score, "improvement_tips": tips}


### backend/resume_parser.py
from pdfminer.high_level import extract_text
from docx import Document
import io

def parse_resume(content):
    try:
        return extract_text(io.BytesIO(content))
    except:
        doc = Document(io.BytesIO(content))
        return "\n".join([para.text for para in doc.paragraphs])

def parse_job_description(content):
    try:
        return extract_text(io.BytesIO(content))
    except:
        doc = Document(io.BytesIO(content))
        return "\n".join([para.text for para in doc.paragraphs])